from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


document = Document()

#profile picture
# document.add_picture('', width= Inches(2.0))

# name email and phone.
name = input ('what is your name? ')
speak('Hello' + name + 'how are you today?')
phone_number= input('what is your phone number? ')
email= input ('What is your email address? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)


# about me
document.add_heading ('About Me')
document.add_paragraph(input ('Tell me about yourself. '))


# Work experiences
document.add_heading('Work Experience')
p=document.add_paragraph()
company= input ('Enter Company Name:  ')
from_date = input('From date: ')
to_date = input ('To Date: ')

p.add_run(company + ' ').bold= True
p.add_run(from_date + '-' + to_date + ' \n' ).italic = True

experience_details= input ('Describe your Experience at ' + company + ' ' )
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences= input ('Do you have more experiences? Yes/no ')
    if has_more_experiences.lower() == 'yes':
        p=document.add_paragraph()
        company= input ('Enter Company Name:  ')
        from_date = input('From date: ')
        to_date = input ('To Date: ')

        p.add_run(company + ' ').bold= True
        p.add_run(from_date + '-' + to_date + ' \n' ).italic = True

        experience_details= input ('Describe your Experience at ' + company + ' ' )
        p.add_run(experience_details)
    else:
        break

# skills
document.add_heading('My Skills')
skill= input ('Enter your skill:  ')
p = document.add_paragraph(skill)
p.style= 'List Bullet'

while True:
    any_more_skill = input('any other skill you have: yes/no  ')
    if any_more_skill.lower() == 'yes':
        skill= input ('Enter your skill:  ')
        p = document.add_paragraph(skill)
        p.style= 'List Bullet'
    else:
        break


# Footer
section = document.sections[0]
footer = section.footer
p= footer.paragraphs[0]
p.text= "cv generated using python docx"

document.save('cv.docx')
